import os, uuid, re
import mimetypes
from pathlib import Path

from flask import current_app, send_file
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename
from werkzeug.exceptions import NotFound

from app.Helpers.registry import register
from app.Util.file import FileUtil
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches


@register("File", repo="File")
class FileService:
    def __init__(self, repository):
        self.repo = repository.get("File")

    def get_by_id(self, file_id: str):
        file = self.repo.get_by_id(file_id)
        if not file:
            return None
        return file

    def get(self, filters=None):
        return self.repo.get(filters)

    def delete(self, file_id: str):
        file_record = self.get_by_id(file_id)
        if not file_record:
            raise ValueError("File not found")

        file_path = getattr(file_record, "file_path", None)
        if file_path and os.path.exists(file_path):
            os.remove(file_path)

        return self.repo.delete(file_record)


    def download_file(self, file_id: str):
        file_record = self.get_by_id(file_id)
        if not file_record:
            raise NotFound("File not found")

        path = FileUtil.get_path(file_record.file_path)
        if path is None or not path.exists():
            raise NotFound("File not found on server")

        return send_file(
            path,
            as_attachment=True,
            download_name=file_record.filename,
            mimetype=file_record.mime_type
        )


    def upload_file(self, file: FileStorage, uploader_id: int):
        if not file or not file.filename:
            raise ValueError("No file provided")

        file.stream.seek(0, os.SEEK_END)
        size = file.stream.tell()
        file.stream.seek(0)
        if size > FileUtil.MAX_FILE_SIZE:
            raise ValueError("Exceeded file size limit")

        mime = file.content_type or mimetypes.guess_type(file.filename)[0] or "application/octet-stream"
        if not FileUtil.allowed_mime(mime):
            raise ValueError("File mime type not allowed")

        if not FileUtil.allowed_file(file.filename):
            raise ValueError("File extension not allowed")

        filename = secure_filename(file.filename)
        upload_folder = current_app.config['UPLOAD_FOLDER']
        os.makedirs(upload_folder, exist_ok=True)

        full_path = os.path.join(upload_folder, filename)
        file.save(full_path)

        file_data = {
            "id": str(uuid.uuid4()),
            "filename": filename,
            "mime_type": mime,
            "file_path": full_path,
            "file_size": os.path.getsize(full_path),
            "uploader_id": uploader_id
        }

        return self.repo.create(file_data)


    def upload_text(self, text, uploader_id: int):
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        if not text:
            raise ValueError("No text provided")

        # Split into paragraphs
        if isinstance(text, str):
            paragraphs = text.split("\n\n")  # Double newline for paragraph breaks
        elif isinstance(text, list):
            paragraphs = text
        else:
            raise ValueError("Input must be a string or a list of paragraphs")

        doc = Document()

        for para_text in paragraphs:
            if not para_text.strip():
                doc.add_paragraph("")
                continue

            # Split by single newlines for lines within paragraph
            lines = para_text.split("\n")
            
            for line in lines:
                if not line.strip():
                    doc.add_paragraph("")
                    continue

                indent = FileUtil.count_leading_spaces(line)
                content = line.lstrip()

                # Detect if text is Arabic
                is_arabic = FileUtil.is_arabic(content)
                

                # Create paragraph and run
                p = doc.add_paragraph()
                run = p.add_run()

                if is_arabic:
                    # Use original text WITHOUT reshaping for Word
                    # Word handles Arabic rendering natively
                    run.text = content
                    
                    # Set font to one that supports Arabic
                    run.font.name = 'Arial'
                    
                    # Also set complex script font
                    rPr = run._element.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(qn('w:ascii'), 'Arial')
                    rFonts.set(qn('w:hAnsi'), 'Arial')
                    rFonts.set(qn('w:cs'), 'Arial')
                    
                    # Enable RTL (right-to-left) at paragraph level
                    pPr = p._element.get_or_add_pPr()
                    
                    # Set bidi (bidirectional) property
                    bidi_elem = OxmlElement('w:bidi')
                    bidi_elem.set(qn('w:val'), '1')
                    pPr.append(bidi_elem)
                    
                    # Set text direction to RTL
                    textDirection = OxmlElement('w:textDirection')
                    textDirection.set(qn('w:val'), 'rl')
                    pPr.append(textDirection)
                    
                    
                    # Set RTL property at run level too
                    rtl_elem = OxmlElement('w:rtl')
                    rPr.append(rtl_elem)
                    
                    print(f"Set as Arabic with RTL")
                else:
                    run.text = content
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    
                # Apply indentation
                if indent > 0:
                    p.paragraph_format.left_indent = Pt(indent * 2)

        # Save DOCX
        upload_folder = current_app.config['UPLOAD_FOLDER']
        os.makedirs(upload_folder, exist_ok=True)
        file_id = str(uuid.uuid4())
        filename = f"{file_id}.docx"
        full_path = os.path.join(upload_folder, filename)
        
        print(f"\n--- Saving document ---")
        print(f"Path: {full_path}")
        
        doc.save(full_path)

        file_data = {
            "id": file_id,
            "filename": filename,
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "file_path": full_path,
            "file_size": os.path.getsize(full_path),
            "uploader_id": uploader_id
        }

        return self.repo.create(file_data)